"""
build_docx.py — render assembled state + synthesized prose into a Word report.

Why python-docx (not docx-js)
-----------------------------
The docx skill recommends docx-js for one-off document creation. This is not a
one-off: it's a node inside a Python LangGraph pipeline that runs on Lilly
infra. python-docx keeps the whole pipeline in one language with no Node
dependency at runtime, which is the right call for programmatic, repeated
generation. The skill's docx-js gotchas (US-Letter sizing, table dual-widths,
CLEAR shading, no literal bullets) still apply conceptually and are honored
here in python-docx terms.

What is rendered from where — every table/figure traces to state
----------------------------------------------------------------
  Title / exec summary / per-COI prose / recommendations  ← report_sections (synthesize.py)
  Device comparison tables (grouped by COI, ordered by tier) ← devices (state)
  Gap log (grouped by severity, blocking first)            ← consolidated gaps (state)
  PRISMA + tier + source figures                           ← figures (corpus_stats.py)
  Methodology (eligibility criteria, verbatim)             ← criteria.py
  Limitations (run_notes + verify orphans + skipped figs)  ← run_notes + verify_report + stats
  Citation appendix (every cited id → its record)          ← citations_index (state)

The prose sections are the ONLY ones that depend on the (possibly degraded)
synthesis step; every hard table/figure comes straight from state, so a
degraded synthesis yields thinner narrative over identical data.

Deterministic. No LLM.
"""
from __future__ import annotations

import logging
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from state import COIEvidence, DeviceRow, Gap, Record


log = logging.getLogger(__name__)

# US Letter, matching the docx skill's guidance (docx-js DXA → python-docx Inches).
_LETTER_W = Inches(8.5)
_LETTER_H = Inches(11)
_ACCENT = RGBColor(0x2F, 0x5C, 0x8A)
_TIER_ORDER = ["Tier 1", "Tier 2", "Tier 3", "Tier 4", "Unclassified"]
_SEV_HEAD = {"blocking": "Blocking", "notable": "Notable", "acknowledged": "Acknowledged"}


# ─────────────────────────────────────────────────────────────────────────────
# Low-level helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_letter(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = _LETTER_W
        section.page_height = _LETTER_H
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)


def _shade_cell(cell, hex_fill: str) -> None:
    """Apply solid-ish shading via CLEAR fill (docx skill: never SOLID)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _cell_text(cell, text: str, *, bold: bool = False, white: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _add_toc(doc: Document) -> None:
    """Insert a Word field-code TOC that populates on open/update."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click and Update Field to populate the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar, instr, fldChar2, t, fldChar3):
        run._r.append(el)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _ACCENT


# ─────────────────────────────────────────────────────────────────────────────
# Section renderers
# ─────────────────────────────────────────────────────────────────────────────

def _render_title(doc: Document, sections: dict, indication: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(sections.get("title", f"DHT Landscape Review — {indication}"))
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = _ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Digital Health Technology landscape assessment · organized research output")
    sr.italic = True
    sr.font.size = Pt(11)

    mode = sections.get("synthesis_mode", "deterministic")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Narrative generated in "
        + {"llm": "analytical-model", "deterministic": "deterministic (template)",
           "deterministic_fallback": "deterministic-fallback (LLM synthesis failed)"}
           .get(mode, mode)
        + " mode."
    )
    nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()


def _render_exec_summary(doc: Document, sections: dict) -> None:
    _heading(doc, "Executive Summary", 1)
    doc.add_paragraph(sections.get("executive_summary", ""))


def _render_device_table(doc: Document, coi: str, rows: list[DeviceRow]) -> None:
    _heading(doc, coi.replace("_", " ").title(), 2)

    # narrative paragraphs for this COI (from synthesis) are added by caller;
    # this function renders only the table.
    cols = ["Device", "Manufacturer", "Tier", "Composite", "Bias",
            "Longitudinal", "Reg. clearance", "Citations"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        _cell_text(hdr[i], c, bold=True, white=True)
        _shade_cell(hdr[i], "2F5C8A")

    for r in rows:
        v3 = r.v3_evidence or {}
        cells = table.add_row().cells
        _cell_text(cells[0], r.device, bold=True)
        _cell_text(cells[1], r.manufacturer)
        _cell_text(cells[2], v3.get("tier", "—"))
        _cell_text(cells[3], v3.get("composite_evidence", "—"))
        _cell_text(cells[4], v3.get("bias_rating", "—"))
        _cell_text(cells[5], v3.get("longitudinal_readiness", "—"))
        _cell_text(cells[6], r.regulatory_clearance or "None found")
        _cell_text(cells[7], ", ".join(r.evidence_citations) or "—")
        # subtle tier shading on the tier cell
        tier_fill = {"Tier 1": "DCE6F1", "Tier 2": "E8EEF5",
                     "Tier 3": "F2F5F9", "Tier 4": "F7F9FB"}.get(v3.get("tier"), None)
        if tier_fill:
            _shade_cell(cells[2], tier_fill)

    # limitations as a compact note under the table
    for r in rows:
        if r.limitations:
            p = doc.add_paragraph()
            run = p.add_run(f"{r.device} — limitations: ")
            run.bold = True
            run.font.size = Pt(8)
            note = p.add_run(r.limitations)
            note.font.size = Pt(8)
            note.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _render_coi_sections(
    doc: Document, sections: dict, devices_by_coi: "OrderedDict[str, list[DeviceRow]]"
) -> None:
    _heading(doc, "Per-COI Assessment", 1)
    prose_by_coi = {s["coi"]: s for s in sections.get("coi_sections", [])}

    for coi, rows in devices_by_coi.items():
        prose = prose_by_coi.get(coi)
        # heading + table
        _render_device_table(doc, coi, rows)
        # narrative below the table (interpretation)
        if prose:
            for para in prose.get("paragraphs", []):
                doc.add_paragraph(para)
        doc.add_paragraph()


def _render_gap_log(doc: Document, gaps_by_sev: "OrderedDict[str, list[Gap]]") -> None:
    _heading(doc, "Gap Log", 1)
    if not gaps_by_sev:
        doc.add_paragraph("No gaps were recorded.")
        return
    doc.add_paragraph(
        "Gaps are consolidated across COIs: a single row's 'Affects' column "
        "lists every COI the gap touches. Blocking gaps must be resolved before "
        "any primary-endpoint use."
    )
    for sev, gaps in gaps_by_sev.items():
        _heading(doc, f"{_SEV_HEAD.get(sev, sev.title())} ({len(gaps)})", 2)
        cols = ["ID", "Category", "Description", "Affects", "Action"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, c in enumerate(cols):
            _cell_text(hdr[i], c, bold=True, white=True)
            _shade_cell(hdr[i], "3E7CB1")
        for g in gaps:
            cells = table.add_row().cells
            _cell_text(cells[0], g.gap_id)
            _cell_text(cells[1], g.category.replace("_", " "))
            _cell_text(cells[2], g.description)
            _cell_text(cells[3], ", ".join(c.replace("_", " ") for c in g.affected_cois))
            _cell_text(cells[4], g.action)
        doc.add_paragraph()


def _render_recommendations(doc: Document, sections: dict) -> None:
    recs = sections.get("recommendations", [])
    if not recs:
        return
    _heading(doc, "Recommendations", 1)
    for rec in recs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rec)


def _render_figures(doc: Document, figures: "OrderedDict[str, str]") -> None:
    if not figures:
        return
    _heading(doc, "Corpus Overview", 1)
    for path, caption in figures.items():
        if not Path(path).exists():
            continue
        doc.add_picture(path, width=Inches(6.0))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        doc.add_paragraph()


def _render_methodology(doc: Document, cois: list[str], criteria_render) -> None:
    _heading(doc, "Methodology", 1)
    doc.add_paragraph(
        "Records were identified from ClinicalTrials.gov, PubMed, Europe PMC, "
        "and Semantic Scholar, de-duplicated by citation ID, screened against "
        "per-COI eligibility criteria, re-adjudicated at eligibility, then scored "
        "on the dht-landscape-scout 12-criterion rubric. Every cited record "
        "resolves in the corpus appendix. The eligibility criteria applied to "
        "each COI are reproduced below verbatim so the screen is reproducible."
    )
    for coi in cois:
        try:
            block = criteria_render(coi)
        except Exception:  # noqa: BLE001 — a synthesized-criteria COI may not be in criteria.py
            block = f"# Eligibility criteria for COI: {coi}\n(Criteria not available in criteria.py — see Limitations.)"
        _heading(doc, coi.replace("_", " ").title(), 2)
        for line in block.splitlines():
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(8)


def _render_limitations(
    doc: Document, run_notes: list[str], verify_report: dict, stats: dict
) -> None:
    _heading(doc, "Limitations and Disclosures", 1)

    orphans = (verify_report or {}).get("orphan_citations", [])
    precision = (verify_report or {}).get("citation_precision", 1.0)
    doc.add_paragraph(
        f"Citation integrity: {(verify_report or {}).get('n_citations_verified', 0)} "
        f"of {(verify_report or {}).get('n_citations_total', 0)} cited tokens "
        f"resolved in the frozen corpus (precision {precision:.2f})."
    )
    if orphans:
        p = doc.add_paragraph()
        p.add_run("Unresolved citations (removed from prose, disclosed here): ").bold = True
        p.add_run(", ".join(orphans))

    skipped = stats.get("skipped_figures", [])
    if skipped:
        doc.add_paragraph("Figures not rendered (empty distributions): " + "; ".join(skipped) + ".")

    for note in (run_notes or []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(note)


def _render_citation_appendix(
    doc: Document, citations_index: dict[str, Record], cited_ids: set[str]
) -> None:
    _heading(doc, "Appendix — Cited Sources", 1)
    used = sorted(cited_ids)
    if not used:
        doc.add_paragraph("No sources were cited in the scored output.")
        return
    for cid in used:
        rec = citations_index.get(cid)
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{cid}")
        run.bold = True
        if rec is not None:
            bits = [rec.title or "(untitled)"]
            if rec.year:
                bits.append(str(rec.year))
            if rec.journal:
                bits.append(rec.journal)
            if rec.doi:
                bits.append(f"doi:{rec.doi}")
            p.add_run(" — " + " · ".join(bits)).font.size = Pt(9)
        else:
            p.add_run(" — (record not present in index)").font.size = Pt(9)


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

def _devices_by_coi(devices: list[DeviceRow]) -> "OrderedDict[str, list[DeviceRow]]":
    out: "OrderedDict[str, list[DeviceRow]]" = OrderedDict()
    for d in devices:
        out.setdefault(d.coi, []).append(d)
    for coi in out:
        out[coi].sort(key=lambda d: (
            _TIER_ORDER.index((d.v3_evidence or {}).get("tier", "Unclassified"))
            if (d.v3_evidence or {}).get("tier", "Unclassified") in _TIER_ORDER else 99,
            d.device,
        ))
    return out


def _cited_ids(devices, gaps, evidence) -> set[str]:
    ids: set[str] = set()
    for d in devices:
        ids.update(d.evidence_citations or [])
    for g in gaps:
        ids.update(g.supporting_citations or [])
    for e in evidence:
        ids.update(e.key_citations or [])
    return ids


def build_docx(
    *,
    indication: str,
    cois: list[str],
    report_sections: dict,
    devices: list[DeviceRow],
    gaps_by_severity: "OrderedDict[str, list[Gap]]",
    figures: "OrderedDict[str, str]",
    citations_index: dict[str, Record],
    verify_report: dict,
    corpus_stats: dict,
    criteria_render,                 # Callable[[str], str] — criteria.render_methodology_block
    out_path: Path,
    run_notes: list[str] | None = None,
) -> Path:
    """Render the full report. Returns the output path."""
    doc = Document()
    _set_letter(doc)

    _render_title(doc, report_sections, indication)

    _heading(doc, "Contents", 1)
    _add_toc(doc)
    doc.add_page_break()

    _render_exec_summary(doc, report_sections)
    _render_figures(doc, figures)

    devices_by_coi = _devices_by_coi(devices)
    _render_coi_sections(doc, report_sections, devices_by_coi)

    _render_recommendations(doc, report_sections)

    # flatten gaps for the appendix cited-id set
    all_gaps = [g for gs in gaps_by_severity.values() for g in gs]
    _render_gap_log(doc, gaps_by_severity)

    _render_methodology(doc, cois, criteria_render)
    _render_limitations(doc, run_notes, verify_report, corpus_stats)

    # evidence objects aren't passed separately here; cited ids come from
    # devices + gaps (evidence key_citations are a subset surfaced in prose).
    cited = _cited_ids(devices, all_gaps, [])
    _render_citation_appendix(doc, citations_index, cited)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    log.info("build_docx: wrote %s", out_path)
    return out_path


__all__ = ["build_docx"]